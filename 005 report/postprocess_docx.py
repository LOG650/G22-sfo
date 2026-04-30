"""Post-process DOCX fra pandoc for å unngå at overskrifter og figurer brytes over sider.

Setter:
- Heading 1-6: keep_with_next=True (ikke alene nederst på side) + keep_together=True
- Bilder (paragrafer som inneholder en `<w:drawing>`): keep_with_next=True
- Figurtekster (kursiv-paragrafer rett etter bilde): keep_together=True
- Alle paragrafer: widow_control=True (default Word-oppførsel, eksplisitt)

Brukes etter pandoc-generering:
    python3 005\\ report/postprocess_docx.py
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
DOCX = ROOT / "rapport.docx"


def set_keep_props(p, keep_with_next=False, keep_together=False, page_break_before=False):
    """Set Word paragraph properties for page-break behavior."""
    pf = p.paragraph_format
    if keep_with_next:
        pf.keep_with_next = True
    if keep_together:
        pf.keep_together = True
    if page_break_before:
        pf.page_break_before = True
    # widow_control eksplisitt på (Word default, men sikrer det)
    pf.widow_control = True


def is_image_paragraph(p) -> bool:
    """Returnerer True hvis paragrafen inneholder en innebygd figur."""
    return p._p.find(qn("w:r") + "/" + qn("w:drawing")) is not None or \
           bool(p._p.findall(".//" + qn("w:drawing")))


def is_caption_paragraph(p) -> bool:
    """Heuristikk: figurtekst-paragraf har 'Figur ' eller 'Tabell ' som første ord, eller er kursiv."""
    text = p.text.strip()
    if not text:
        return False
    return (text.startswith("Figur ") or text.startswith("Tabell ")) or \
           any(r.italic for r in p.runs if r.italic is not None and r.italic)


def remove_numbering(p) -> None:
    """Fjerner auto-numerering (numPr) som HiMolde-malen påfører Heading-stiler.

    Pandoc respekterer mal-stilenes numbering scheme, men malen har en list-numerator
    som teller feil (gir '1.11.27.2 ...' istedenfor '7.2'). Vi fjerner numPr fra
    paragrafens egen pPr så Word stoler på pandoc-rendret nummerering i selve teksten.
    """
    pPr = p._p.find(qn("w:pPr"))
    if pPr is None:
        return
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)


def strip_numbering_from_styles(doc) -> int:
    """Fjern auto-numerering fra Heading-stilene i mal-en (numPr på styles-nivå)."""
    styles = doc.part.styles_xml = doc.styles.element
    n = 0
    for style in doc.styles:
        name = (style.name or "").lower()
        if not (name.startswith("heading") or name.startswith("overskrift") or name.startswith("title")):
            continue
        # style._element er styles XML for denne stilen
        elem = style.element
        pPr = elem.find(qn("w:pPr"))
        if pPr is not None:
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                pPr.remove(numPr)
                n += 1
    return n


def main() -> None:
    doc = Document(str(DOCX))
    n_styles_cleaned = strip_numbering_from_styles(doc)
    n_headings = 0
    n_images = 0
    n_captions = 0
    prev = None
    for p in doc.paragraphs:
        style = (p.style.name or "").lower() if p.style else ""

        # Overskrifter: keep_with_next + keep_together + fjern auto-numerering
        if style.startswith("heading") or style.startswith("overskrift") or style.startswith("title"):
            set_keep_props(p, keep_with_next=True, keep_together=True)
            remove_numbering(p)
            n_headings += 1
            prev = p
            continue

        # Bildeparagraf: keep_with_next så figurtekst følger med
        if is_image_paragraph(p):
            set_keep_props(p, keep_with_next=True, keep_together=True)
            n_images += 1
            prev = p
            continue

        # Figurtekst rett etter bilde: keep_together (ikke split over side)
        if is_caption_paragraph(p):
            set_keep_props(p, keep_together=True)
            n_captions += 1
            prev = p
            continue

        # Vanlige paragrafer: widow_control eksplisitt
        set_keep_props(p)
        prev = p

    # Tabeller: alle rader keep_together + synlige rammer + autofit til innhold
    from docx.oxml import OxmlElement
    n_tables = 0
    for tbl in doc.tables:
        # Per-rad: cantSplit (radene brytes ikke midt i)
        for row in tbl.rows:
            trPr = row._tr.get_or_add_trPr()
            if trPr.find(qn("w:cantSplit")) is None:
                trPr.append(OxmlElement("w:cantSplit"))

        # Rammer på hele tabellen + alle celler
        tblPr = tbl._tbl.tblPr
        # Fjern eksisterende borders hvis noen, legg til nye
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")  # 0.5 pt
            b.set(qn("w:color"), "808080")  # grå, mer subtil enn svart
            borders.append(b)
        tblPr.append(borders)

        # Autofit til innhold — Word justerer kolonnebredder etter tekst
        tblLayout = tblPr.find(qn("w:tblLayout"))
        if tblLayout is None:
            tblLayout = OxmlElement("w:tblLayout")
            tblPr.append(tblLayout)
        tblLayout.set(qn("w:type"), "autofit")

        # Fjern eksisterende tblW (preferred width) som låser bredde til en pct
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is not None:
            tblW.set(qn("w:w"), "0")
            tblW.set(qn("w:type"), "auto")

        n_tables += 1

    doc.save(str(DOCX))
    print(f"Post-prosessert {DOCX.name}:")
    print(f"  Overskrift-stiler renset for numPr: {n_styles_cleaned}")
    print(f"  Overskrifter: {n_headings}")
    print(f"  Bilder: {n_images}")
    print(f"  Figurtekster: {n_captions}")
    print(f"  Tabeller: {n_tables}")


if __name__ == "__main__":
    main()
